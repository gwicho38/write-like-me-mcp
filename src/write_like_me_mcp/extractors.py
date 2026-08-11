"""Text extraction from various document formats."""

import logging
import re
from pathlib import Path

import chardet

logger = logging.getLogger(__name__)

#: Extensions treated as Markdown, and so stripped of non-prose structure.
MARKDOWN_EXTENSIONS: frozenset[str] = frozenset({".md", ".markdown"})

#: Leading YAML/TOML frontmatter block, only when it opens the document.
_FRONTMATTER_RE = re.compile(r"\A(?:---|\+\+\+)\r?\n.*?\r?\n(?:---|\+\+\+)\s*?(?:\r?\n|\Z)", re.DOTALL)

#: Fenced code blocks (``` or ~~~), including the info string and body.
_FENCED_CODE_RE = re.compile(r"^[ \t]*(`{3,}|~{3,}).*?(?:^[ \t]*\1[ \t]*$|\Z)", re.DOTALL | re.MULTILINE)


def strip_markdown_structure(text: str) -> str:
    """Remove non-prose Markdown scaffolding so only authored sentences remain.

    Frontmatter keys and code identifiers are metadata, not writing. Leaving
    them in lets ``tags: [...]`` lines count as sentences, which skews the
    sentence-length and passive-voice metrics and lets metadata bigrams
    outrank real phrases in the signature-phrase ranking.

    Args:
        text: Raw Markdown source.

    Returns:
        The text with leading frontmatter and fenced code blocks removed.
    """
    without_frontmatter = _FRONTMATTER_RE.sub("", text, count=1)
    return _FENCED_CODE_RE.sub("", without_frontmatter)


def extract_text(file_path: Path) -> str | None:
    """Extract text content from a file based on its extension.

    Args:
        file_path: Path to the file to extract text from.

    Returns:
        Extracted text content, or None if extraction fails.
    """
    suffix = file_path.suffix.lower()
    try:
        if suffix in MARKDOWN_EXTENSIONS:
            return strip_markdown_structure(_extract_plain_text(file_path))
        elif suffix == ".txt":
            return _extract_plain_text(file_path)
        elif suffix == ".pdf":
            return _extract_pdf(file_path)
        elif suffix == ".docx":
            return _extract_docx(file_path)
        elif suffix in (".html", ".htm"):
            return _extract_html(file_path)
        else:
            logger.warning("Unsupported file extension: %s for %s", suffix, file_path)
            return None
    except Exception as e:
        logger.error("Failed to extract text from %s: %s", file_path, e)
        return None


def _read_with_encoding(file_path: Path) -> str:
    """Read a file, detecting encoding if needed."""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        raw = file_path.read_bytes()
        detected = chardet.detect(raw)
        encoding = detected.get("encoding", "utf-8") or "utf-8"
        return raw.decode(encoding, errors="replace")


def _extract_plain_text(file_path: Path) -> str:
    """Extract text from plain text or markdown files."""
    return _read_with_encoding(file_path)


def _extract_pdf(file_path: Path) -> str:
    """Extract text from PDF files using PyMuPDF (fitz)."""
    import fitz  # PyMuPDF

    text_parts: list[str] = []
    with fitz.open(str(file_path)) as doc:
        for page_num, page in enumerate(doc, 1):
            page_text = page.get_text("text")
            if page_text.strip():
                text_parts.append(f"[Page {page_num}]\n{page_text}")
    return "\n\n".join(text_parts)


def _extract_docx(file_path: Path) -> str:
    """Extract text from Word .docx files."""
    from docx import Document

    doc = Document(str(file_path))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Preserve heading structure
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "").strip()
                try:
                    hashes = "#" * int(level)
                    parts.append(f"{hashes} {text}")
                except ValueError:
                    parts.append(text)
            else:
                parts.append(text)

    # Also extract text from tables
    for table in doc.tables:
        table_rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(" | ".join(cells))
        if table_rows:
            parts.append("\n".join(table_rows))

    return "\n\n".join(parts)


def _extract_html(file_path: Path) -> str:
    """Extract text from HTML files."""
    from bs4 import BeautifulSoup

    content = _read_with_encoding(file_path)
    soup = BeautifulSoup(content, "html.parser")

    # Remove script and style elements
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()

    text = soup.get_text(separator="\n")
    # Collapse multiple blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


